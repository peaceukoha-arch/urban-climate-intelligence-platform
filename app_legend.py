
# ======================================
# IMPORT LIBRARIES
# ======================================

import zipfile
import os
import tempfile
import ee
import geopandas as gpd
import pandas as pd
import streamlit as st
import matplotlib.pyplot as plt
import folium
import numpy as np


from streamlit_folium import st_folium
from datetime import date
from docx import Document 
from docx.shared import Inches


# ======================================
# INITIALIZE EARTH ENGINE
# ======================================

try:

    ee.Initialize()

except:

    ee.Authenticate()

    ee.Initialize()


# ======================================
# LANDSAT 8 + 9 COLLECTION
# ======================================

landsat8 = ee.ImageCollection(
    'LANDSAT/LC08/C02/T1_L2'
)

landsat9 = ee.ImageCollection(
    'LANDSAT/LC09/C02/T1_L2'
)

landsat = landsat8.merge(
    landsat9
)


# ======================================
# PAGE CONFIG
# ======================================

st.set_page_config(
    page_title="Urban Heat Intelligence",
    layout="wide"
)


# ======================================
# SIDEBAR
# ======================================

st.sidebar.title(
    "Urban Climate Controls"
)

st.sidebar.write(
    "Configure analysis settings"
)


# ======================================
# NAVIGATION
# ======================================


page = st.sidebar.radio(
    "Navigation",
    [
        "Map Analysis",
        "Time-Series Analytics",
        "Map Layout Export",
        "Report Analysis"
    ]
)
# ======================================
# MAP ANALYSIS PAGE
# ======================================

if page == "Map Analysis":

    st.title(
        "Urban Heat Intelligence Platform"
    )

    st.write(
        """
        Upload GeoJSON, KML, or ZIP Shapefile
        to generate:
        - Land Surface Temperature
        - NDVI
        - Hotspot Analysis
        """
    )

    # ==================================
    # DATE RANGE
    # ==================================

    start_date = st.sidebar.date_input(
        "Start Date",
        date(2024, 1, 1)
    )

    end_date = st.sidebar.date_input(
        "End Date",
        date(2024, 12, 31)
    )

    # ==================================
    # BASEMAP
    # ==================================

    basemap = st.sidebar.selectbox(
        "Select Basemap",
        [
            "OpenStreetMap",
            "SATELLITE",
            "HYBRID",
            "TERRAIN"
        ]
    )

    # ==================================
    # LAYER TOGGLES
    # ==================================

    show_lst = st.sidebar.checkbox(
        "Show LST",
        value=True
    )

    show_ndvi = st.sidebar.checkbox(
        "Show NDVI",
        value=True
    )

    show_hotspot = st.sidebar.checkbox(
        "Show Hotspot",
        value=True
    )

    # ==================================
    # FILE UPLOAD
    # ==================================

    uploaded_file = st.sidebar.file_uploader(
        "Upload Boundary",
        type=["geojson", "kml", "zip"]
    )

    # ==================================
    # PROCESS FILE
    # ==================================

    if uploaded_file is not None:

        file_extension = os.path.splitext(
            uploaded_file.name
        )[1]

        temp_file = tempfile.NamedTemporaryFile(
            delete=False,
            suffix=file_extension
        )

        temp_file.write(
            uploaded_file.read()
        )

        temp_file.close()

        # ==============================
        # READ FILES
        # ==============================

        if uploaded_file.name.endswith(
            ".geojson"
        ):

            gdf = gpd.read_file(
                temp_file.name
            )

        elif uploaded_file.name.endswith(
            ".kml"
        ):

            gdf = gpd.read_file(
                temp_file.name,
                driver="KML"
            )

        elif uploaded_file.name.endswith(
            ".zip"
        ):

            extract_path = tempfile.mkdtemp()

            with zipfile.ZipFile(
                temp_file.name,
                "r"
            ) as zip_ref:

                zip_ref.extractall(
                    extract_path
                )

            shp_file = None

            for root, dirs, files in os.walk(
                extract_path
            ):

                for file in files:

                    if file.endswith(".shp"):

                        shp_file = os.path.join(
                            root,
                            file
                        )

            gdf = gpd.read_file(
                shp_file
            )

        st.success(
            "Boundary uploaded successfully!"
        )

        # ==============================
        # EARTH ENGINE GEOMETRY
        # ==============================

        geojson = gdf.__geo_interface__

        ee_geometry = ee.FeatureCollection(
            geojson
        )

        # ==============================
        # CLOUD MASK
        # ==============================

        def mask_clouds(image):

            qa = image.select(
                'QA_PIXEL'
            )

            cloud = qa.bitwiseAnd(
                1 << 3
            ).eq(0)

            cloud_shadow = qa.bitwiseAnd(
                1 << 4
            ).eq(0)

            mask = cloud.And(
                cloud_shadow
            )

            return image.updateMask(mask)

        # ==============================
        # CALCULATE INDICES
        # ==============================

        def calculate_indices(image):

            thermal = image.select(
                'ST_B10'
            ).multiply(
                0.00341802
            ).add(
                149.0
            )

            lst = thermal.subtract(
                273.15
            ).rename(
                'LST'
            )

            ndvi = image.normalizedDifference(
                ['SR_B5', 'SR_B4']
            ).rename(
                'NDVI'
            )

            return image.addBands([
                lst,
                ndvi
            ])

        # ==============================
        # PROCESS COLLECTION
        # ==============================

        processed = landsat \
            .filterBounds(
                ee_geometry
            ) \
            .filterDate(
                str(start_date),
                str(end_date)
            ) \
            .filter(
                ee.Filter.lt(
                    'CLOUD_COVER',
                    20
                )
            ) \
            .map(
                mask_clouds
            ) \
            .map(
                calculate_indices
            )

        # ==============================
        # CREATE COMPOSITES
        # ==============================

        lst_image = processed.select(
            'LST'
        ).median().clip(
            ee_geometry
        )

        lst_stats = lst_image.reduceRegion(
            reducer=ee.Reducer.minMax(),
            geometry=ee_geometry,
            scale=30,
            maxPixels=1e13
        )

        lst_min = lst_stats.getInfo()['LST_min']
        lst_max = lst_stats.getInfo()['LST_max']

        interval = (lst_max - lst_min) / 4

        class1 = lst_min + interval
        class2 = lst_min + (interval * 2)
        class3 = lst_min + (interval * 3)


        ndvi_image = processed.select(
            'NDVI'
        ).median().clip(
            ee_geometry
        )

        # ==============================
        # HOTSPOT CLASSIFICATION
        # ==============================

        hotspot = lst_image.expression(

            "(b('LST') < 25) ? 1" +
            ": (b('LST') < 30) ? 2" +
            ": (b('LST') < 35) ? 3" +
            ": (b('LST') < 40) ? 4" +
            ": 5"

        ).rename(
            'Hotspot'
        ).clip(
            ee_geometry
        )

        # ==============================
        # LST LEGEND COLORS
        # ==============================

        st.sidebar.subheader(
            "LST Legend Colors"
        )

        lst_low = st.sidebar.color_picker(
            "Low",
            "#00FF00"
        )
        
        lst_moderate = st.sidebar.color_picker(
            "Moderate",
            "#FFFF00"
        )
    
        lst_high = st.sidebar.color_picker(
            "High",
            "#FFA500"
        )

        lst_extreme = st.sidebar.color_picker(
            "Extreme",
            "#FF0000"
        )

        # ==============================
        # NDVI LEGEND COLORS
        # ==============================

        st.sidebar.subheader(
            "NDVI Legend Colors"
        )

        ndvi_low = st.sidebar.color_picker(
            "Low Vegetation",
            "#8B4513"
        )

        ndvi_moderate = st.sidebar.color_picker(
            "Moderate Vegetation",
            "#FFFF00"
        )

        ndvi_high = st.sidebar.color_picker(
            "High Vegetation",
            "#008000"
        )

        # ==============================
        # HOTSPOT LEGEND COLORS
        # ==============================

        st.sidebar.subheader(
            "Hotspot Legend Colors"
        )

        hotspot_low = st.sidebar.color_picker(
            "Low Hotspot",
            "#abd9e9"
        )

        hotspot_moderate = st.sidebar.color_picker(
            "Moderate Hotspot",
            "#ffffbf"
        )

        hotspot_high = st.sidebar.color_picker(
            "High Hotspot",
            "#fdae61"
        )

        hotspot_extreme = st.sidebar.color_picker(
            "Extreme Hotspot",
            "#d7191c"
        )

        # ==============================
        # VISUALIZATION PARAMETERS
        # ==============================

        lst_vis = {
            'min': 1,
            'max': 4,
            'palette': [
                str(lst_low).replace("#", ""),
                str(lst_moderate).replace("#", ""),
                str(lst_high).replace("#", ""),
                str(lst_extreme).replace("#", "")
            ]
        }

        ndvi_vis = {
            'min': -1,
            'max': 1,
            'palette': [
                str(ndvi_low).replace("#", ""),
                str(ndvi_moderate).replace("#", ""),
                str(ndvi_high).replace("#", "")
            ]
        }

        hotspot_vis = {
            'min': 1,
            'max': 4,
            'palette': [
                str(hotspot_low).replace("#", ""),
                str(hotspot_moderate).replace("#", ""),
                str(hotspot_high).replace("#", ""),
                str(hotspot_extreme).replace("#", "")
            ]
        }

        # ==============================
        # MAP CENTER
        # ==============================

        centroid = gdf.geometry.centroid.iloc[0]

        map_center = [
        centroid.y,
        centroid.x
        ]
        # ==============================
        # CREATE MAP
        # ==============================

        m = folium.Map(
            location=map_center,
            zoom_start=8,
            control_scale=True
        )

        # ==============================
        # BASEMAPS
        # ==============================

        if basemap == "OpenStreetMap":

            folium.TileLayer(
                'OpenStreetMap',
                name='OpenStreetMap'
            ).add_to(m)

        elif basemap == "SATELLITE":

            folium.TileLayer(
                tiles='https://mt1.google.com/vt/lyrs=s&x={x}&y={y}&z={z}',
                attr='Google',
                name='Satellite'
            ).add_to(m)

        elif basemap == "HYBRID":

            folium.TileLayer(
                tiles='https://mt1.google.com/vt/lyrs=y&x={x}&y={y}&z={z}',
                attr='Google',
                name='Hybrid'
            ).add_to(m)

        elif basemap == "TERRAIN":

            folium.TileLayer(
                tiles='https://mt1.google.com/vt/lyrs=p&x={x}&y={y}&z={z}',
                attr='Google',
                name='Terrain'
            ).add_to(m)

        # ==============================
        # STUDY AREA OUTLINE
        # ==============================

        folium.GeoJson(
            geojson,
            name="Study Area",
            style_function=lambda x: {
                'fillColor': 'none',
                'color': 'black',
                'weight': 2
            }
        ).add_to(m)

        # ==============================
        # CLASSIFIED LST
        # ==============================

        classified_lst = (
            lst_image.gt(lst_min).And(
                lst_image.lte(class1)
            ).multiply(1)
            .add(
                lst_image.gt(class1).And(
                    lst_image.lte(class2)
                ).multiply(2)
            )
            .add(
                lst_image.gt(class2).And(
                    lst_image.lte(class3)
                ).multiply(3)
            )
            .add(
                lst_image.gt(class3).multiply(4)
            )
        )

        
        # ==============================
        # LST LAYER
        # ==============================

        if show_lst:

            map_id = classified_lst.getMapId(
                lst_vis
        )

            folium.TileLayer(
                tiles=map_id['tile_fetcher'].url_format,
                attr='Google Earth Engine',
                name='LST',
                overlay=True,
                control=True
            ).add_to(m)

        # ==============================
        # NDVI LAYER
        # ==============================

        if show_ndvi:

            map_id = ee.Image(
                ndvi_image
            ).getMapId(ndvi_vis)

            folium.TileLayer(
                tiles=map_id['tile_fetcher'].url_format,
                attr='Google Earth Engine',
                name='NDVI',
                overlay=True,
                control=True
            ).add_to(m)

        # ==============================
        # HOTSPOT LAYER
        # ==============================

        if show_hotspot:

            map_id = ee.Image(
                hotspot
            ).getMapId(
                hotspot_vis
            )

            folium.TileLayer(
                tiles=map_id['tile_fetcher'].url_format,
                attr='Google Earth Engine',
                name='Hotspot',
                overlay=True,
                control=True
            ).add_to(m)

        # ==============================
        # LAYER CONTROL
        # ==============================

        folium.LayerControl().add_to(m)

        # ==============================
        # COMBINED LEGENDS
        # ==============================

        legend_html = ""

        # ==============================
        # LST LEGEND
        # ==============================

        
        if show_lst:

            legend_html += f"""

            <div style="
            position: fixed;
            bottom: 20px;
            right: 20px;
            width: 180px;
            height: 150px;
            background-color: white;
            border:2px solid grey;
            z-index:9999;
            font-size:14px;
            padding: 10px;
            border-radius:8px;
            ">

            <b>LST (°C)</b><br><br>

            <i style="background:{lst_low};
            width:15px;
            height:15px;
            float:left;
            margin-right:8px;"></i>
            Low<br>

            <i style="background:{lst_moderate};
            width:15px;
            height:15px;
            float:left;
            margin-right:8px;"></i>
            Moderate<br>

            <i style="background:{lst_high};
            width:15px;
            height:15px;
            float:left;
            margin-right:8px;"></i>
            High<br>

            <i style="background:{lst_extreme};
            width:15px;
            height:15px;
            float:left;
            margin-right:8px;"></i>
            Extreme

            </div>

            """


        # ==============================
        # NDVI LEGEND
        # ==============================

        if show_ndvi:

            legend_html += f"""

            <div style="
            position: fixed;
            bottom: 20px;
            right: 20px;
            width: 180px;
            height: 130px;
            background-color: white;
            border:2px solid grey;
            z-index:9999;
            font-size:14px;
            padding: 10px;
            border-radius:8px;
            ">

            <b>NDVI</b><br><br>

            <i style="background:{ndvi_low};
            width:15px;
            height:15px;
            float:left;
            margin-right:8px;"></i>
            Low Vegetation<br>

            <i style="background:{ndvi_moderate};
            width:15px;
            height:15px;
            float:left;
            margin-right:8px;"></i>
            Moderate<br>

            <i style="background:{ndvi_high};
            width:15px;
            height:15px;
            float:left;
            margin-right:8px;"></i>
            High Vegetation

            </div>

            """

        # ==============================
        # HOTSPOT LEGEND
        # ==============================

        if show_hotspot:

            legend_html += f"""

            <div style="
            position: fixed;
            bottom: 20px;
            right: 20px;
            width: 180px;
            height: 170px;
            background-color: white;
            border:2px solid grey;
            z-index:9999;
            font-size:14px;
            padding: 10px;
            border-radius:8px;
            ">

            <b>Heat Hotspots</b><br><br>

            <i style="background:{hotspot_low};
            width:15px;
            height:15px;
            float:left;
            margin-right:8px;"></i>
            Low<br>

            <i style="background:{hotspot_moderate};
            width:15px;
            height:15px;
            float:left;
            margin-right:8px;"></i>
            Moderate<br>

            <i style="background:{hotspot_high};
            width:15px;
            height:15px;
            float:left;
            margin-right:8px;"></i>
            High<br>

            <i style="background:{hotspot_extreme};
            width:15px;
            height:15px;
            float:left;
            margin-right:8px;"></i>
            Extreme

            </div>

            """

        m.get_root().html.add_child(
            folium.Element(
                legend_html
            )
        )

        # ==============================
        # DISPLAY MAP
        # ==============================

        if uploaded_file is not None:

            st_folium(
                m,
                width=1200,
                height=700,
                key=(
                    f"{lst_low}_{lst_moderate}_{lst_high}_{lst_extreme}_"
                    f"{ndvi_low}_{ndvi_moderate}_{ndvi_high}_"
                    f"{hotspot_low}_{hotspot_moderate}_{hotspot_high}_{hotspot_extreme}"
                )
            )


        st.markdown("---")

        st.subheader(
            "Analysis Summary"
        )

        col1, col2, col3 = st.columns(3)

        # ==============================
        # LST SUMMARY
        # ==============================

        if show_lst:

            with col1:

                st.metric(
                    "Minimum LST (°C)",
                    f"{lst_min:.2f}"
                )

                st.metric(
                    "Maximum LST (°C)",
                    f"{lst_max:.2f}"
                )

# ==============================
# NDVI SUMMARY
# ==============================

        if show_ndvi:

            ndvi_stats = ndvi_image.reduceRegion(
                reducer=ee.Reducer.mean(),
                geometry=ee_geometry,
                scale=30,
                maxPixels=1e13
            )

            ndvi_mean = ndvi_stats.getInfo().get(
                "NDVI",
                0
            )

            with col2:

                st.metric(
                    "Mean NDVI",
                    f"{ndvi_mean:.2f}"
            )

# ==============================
# HOTSPOT SUMMARY
# ==============================

        if show_hotspot:

            hotspot_stats = hotspot.reduceRegion(
                reducer=ee.Reducer.mean(),
                geometry=ee_geometry,
                scale=30,
                maxPixels=1e13
            )

            hotspot_mean = hotspot_stats.getInfo().get(
                "Hotspot",
                0
            )

            st.button(
                "Download Map PNG",
                key="download_map_png"
            )
# ======================================
# TIMESERIES PAGE
# ======================================

if page == "Time-Series Analytics":

    st.title(
        "Climate Time-Series Analytics"
    )

    st.write(
        """
        Analyze:
        - Yearly Trends
        - Monthly Trends
        - Correlation Analysis
        """
    )

    # ==================================
    # TIMESERIES TYPE
    # ==================================

    timeseries_type = st.sidebar.radio(
        "Time-Series Type",
        [
            "Yearly",
            "Monthly"
        ]
    )

    # ==================================
    # ANALYSIS TYPE
    # ==================================

    analysis_type = st.sidebar.selectbox(
        "Analysis Type",
        [
            "LST Trend",
            "NDVI Trend",
            "LST-NDVI Correlation"
        ]
    )

    # ==================================
    # YEARLY SETTINGS
    # ==================================

    if timeseries_type == "Yearly":

        start_year = st.sidebar.selectbox(
            "Start Year",
            list(range(2015, 2026)),
            index=0
        )

        end_year = st.sidebar.selectbox(
            "End Year",
            list(range(2015, 2026)),
            index=10
        )

    # ==================================
    # MONTHLY SETTINGS
    # ==================================

    elif timeseries_type == "Monthly":

        start_month = st.sidebar.date_input(
            "Start Month",
            date(2015, 1, 1)
        )

        end_month = st.sidebar.date_input(
            "End Month",
            date.today()
        )

    # ==================================
    # FILE UPLOAD
    # ==================================

    uploaded_file = st.sidebar.file_uploader(
        "Upload Boundary",
        type=["geojson", "kml", "zip"],
        key="timeseries_upload"
    )

    # ==================================
    # PROCESS FILE
    # ==================================

    if uploaded_file is not None:

        file_extension = os.path.splitext(
            uploaded_file.name
        )[1]

        temp_file = tempfile.NamedTemporaryFile(
            delete=False,
            suffix=file_extension
        )

        temp_file.write(
            uploaded_file.read()
        )

        temp_file.close()

        # ==============================
        # READ FILES
        # ==============================

        if uploaded_file.name.endswith(".geojson"):

            gdf = gpd.read_file(
                temp_file.name
            )

        elif uploaded_file.name.endswith(".kml"):

            gdf = gpd.read_file(
                temp_file.name,
                driver="KML"
            )

        elif uploaded_file.name.endswith(".zip"):

            extract_path = tempfile.mkdtemp()

            with zipfile.ZipFile(
                temp_file.name,
                "r"
            ) as zip_ref:

                zip_ref.extractall(
                    extract_path
                )

            shp_file = None

            for root, dirs, files in os.walk(
                extract_path
            ):

                for file in files:

                    if file.endswith(".shp"):

                        shp_file = os.path.join(
                            root,
                            file
                        )

            gdf = gpd.read_file(
                shp_file
            )

        st.success(
            "Boundary uploaded successfully!"
        )

        # ==============================
        # EARTH ENGINE GEOMETRY
        # ==============================

        geojson = gdf.__geo_interface__

        ee_geometry = ee.FeatureCollection(
            geojson
        )

        # ==============================
        # CLOUD MASK
        # ==============================

        def mask_clouds(image):

            qa = image.select(
                'QA_PIXEL'
            )

            cloud = qa.bitwiseAnd(
                1 << 3
            ).eq(0)

            cloud_shadow = qa.bitwiseAnd(
                1 << 4
            ).eq(0)

            mask = cloud.And(
                cloud_shadow
            )

            return image.updateMask(mask)

        # ==============================
        # CALCULATE INDICES
        # ==============================

        def calculate_indices(image):

            thermal = image.select(
                'ST_B10'
            ).multiply(
                0.00341802
            ).add(
                149.0
            )

            lst = thermal.subtract(
                273.15
            ).rename(
                'LST'
            )

            ndvi = image.normalizedDifference(
                ['SR_B5', 'SR_B4']
            ).rename(
                'NDVI'
            )

            return image.addBands([
                lst,
                ndvi
            ])

        # ==============================
        # PERIODS
        # ==============================

        if timeseries_type == "Yearly":

            periods = list(
                range(start_year, end_year + 1)
            )

            x_values = periods

        else:

            monthly_dates = pd.date_range(
                start=start_month,
                end=end_month,
                freq='MS'
            )

            periods = monthly_dates

            x_values = [
                d.strftime("%b-%Y")
                for d in monthly_dates
            ]

        lst_values = []

        ndvi_values = []

        # ==============================
        # ANALYSIS LOOP
        # ==============================

        with st.spinner(
            "Generating analysis..."
        ):

            for period in periods:

                if timeseries_type == "Yearly":

                    start = f"{period}-01-01"

                    end = f"{period}-12-31"

                else:

                    start = period.strftime(
                        "%Y-%m-%d"
                    )

                    if period.month == 12:

                        next_month = period.replace(
                            year=period.year + 1,
                            month=1
                        )

                    else:

                        next_month = period.replace(
                            month=period.month + 1
                        )

                    end = next_month.strftime(
                        "%Y-%m-%d"
                    )

                collection = landsat \
                    .filterBounds(
                        ee_geometry
                    ) \
                    .filterDate(
                        start,
                        end
                    ) \
                    .filter(
                        ee.Filter.lt(
                            'CLOUD_COVER',
                            20
                        )
                    ) \
                    .map(
                        mask_clouds
                    ) \
                    .map(
                        calculate_indices
                    )

                # ==========================
                # LST
                # ==========================

                lst_image = collection.select(
                    'LST'
                ).median()

                lst_stats = lst_image.reduceRegion(
                    reducer=ee.Reducer.mean(),
                    geometry=ee_geometry,
                    scale=500,
                    bestEffort=True
                )

                lst_mean = lst_stats.getInfo().get(
                    'LST',
                    None
                )

                if lst_mean is None:

                    if len(lst_values) > 0:

                        lst_mean = round(
                            sum(lst_values) / len(lst_values),
                            2
                        )

                    else:

                        lst_mean = 0

                else:

                    lst_mean = round(
                        lst_mean,
                        2
                    )

                lst_values.append(
                    lst_mean
                )

                # ==========================
                # NDVI
                # ==========================

                ndvi_image = collection.select(
                    'NDVI'
                ).median()

                ndvi_stats = ndvi_image.reduceRegion(
                    reducer=ee.Reducer.mean(),
                    geometry=ee_geometry,
                    scale=500,
                    bestEffort=True
                )

                ndvi_mean = ndvi_stats.getInfo().get(
                    'NDVI',
                    None
                )

                if ndvi_mean is None:

                    if len(ndvi_values) > 0:

                        ndvi_mean = round(
                            sum(ndvi_values) / len(ndvi_values),
                            3
                        )

                    else:

                        ndvi_mean = 0

                else:

                    ndvi_mean = round(
                        ndvi_mean,
                        3
                    )

                ndvi_values.append(
                    ndvi_mean
                )

        # ==============================
        # DATAFRAME
        # ==============================

        df = pd.DataFrame({

            'Period': x_values,

            'Mean_LST': lst_values,

            'Mean_NDVI': ndvi_values

        })

        st.subheader(
            "Time-Series Data"
        )

        st.dataframe(df)

        # ==============================
        # PLOT
        # ==============================

        fig, ax = plt.subplots(
            figsize=(12, 5)
        )

        if analysis_type == "LST Trend":

            ax.plot(
                x_values,
                lst_values,
                marker='o'
            )

            ax.set_ylabel(
                "Mean LST (°C)"
            )

        elif analysis_type == "NDVI Trend":

            ax.plot(
                x_values,
                ndvi_values,
                marker='s'
            )

            ax.set_ylabel(
                "Mean NDVI"
            )

        elif analysis_type == "LST-NDVI Correlation":

            correlation = df[
                'Mean_LST'
            ].corr(
                df['Mean_NDVI']
            )

            r2 = correlation ** 2

            col1, col2 = st.columns(2)

            with col1:

                st.metric(
                    "Correlation",
                    round(correlation, 3)
                )

            with col2:

                st.metric(
                    "R²",
                    round(r2, 3)
                )

            ax.scatter(
                ndvi_values,
                lst_values
            )

            ax.set_xlabel(
                "Mean NDVI"
            )

            ax.set_ylabel(
                "Mean LST (°C)"
            )

        ax.grid(True)

        plt.xticks(
            rotation=45
        )

        st.pyplot(fig)

        # ==============================
        # EXPORT PNG
        # ==============================

        chart_path = "timeseries_chart.png"

        fig.savefig(
            chart_path,
            dpi=300,
            bbox_inches='tight'
        )

        with open(
            chart_path,
            "rb"
        ) as file:

            st.download_button(
                label="Download Chart PNG",
                data=file,
                file_name="Climate_Timeseries_Chart.png",
                mime="image/png"
            )

        # ==============================
        # EXPORT EXCEL
        # ==============================

        excel_path = "timeseries_data.xlsx"

        df.to_excel(
            excel_path,
            index=False
        )

        with open(
            excel_path,
            "rb"
        ) as file:

            st.download_button(
                label="Download Excel Data",
                data=file,
                file_name="Climate_Timeseries_Data.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
# ======================================
# MAP LAYOUT EXPORT PAGE
# ======================================      

if page == "Map Layout Export":

    st.title(
        "Cartographic Layout Studio"
    )

    st.write(
        """
        Create professional cartographic layouts
        for reports, presentations and publications.
        """
    )

    # ==============================
    # MAP SETTINGS
    # ==============================

    map_title = st.text_input(
        "Map Title",
        ""
    )

    subtitle = st.text_input(
        "Subtitle (Optional)",
        ""
    )

    map_type = st.selectbox(
        "Map Type",
        [
            "Land Surface Temperature (LST)",
            "NDVI",
            "Heat Hotspot"
        ]
    )

    paper_size = st.selectbox(
        "Paper Size",
        [
            "A4 Portrait",
            "A4 Landscape",
            "A3 Portrait",
            "A3 Landscape"
        ]
    )

    # ==============================
    # CARTOGRAPHIC ELEMENTS
    # ==============================

    st.subheader(
        "Cartographic Elements"
    )

    show_legend = st.checkbox(
        "Show Legend",
        value=True
    )

    show_north_arrow = st.checkbox(
        "Show North Arrow",
        value=True
    )

    show_scale_bar = st.checkbox(
        "Show Scale Bar",
        value=True
    )

    # ==============================
    # PREVIEW
    # ==============================

    st.markdown("---")

    st.subheader(
        "Layout Preview"
    )

    st.markdown(
        f"""
        ## {map_title}

        {subtitle}

        ---
        🗺️ MAP AREA

        (LST / NDVI / Hotspot map will be inserted here)

        ---

        Map Type: {map_type}

        Paper Size: {paper_size}
        """
    )

    if show_legend:
        st.success("Legend Enabled")

    if show_north_arrow:
        st.success("North Arrow Enabled")

    if show_scale_bar:
        st.success("Scale Bar Enabled")

    st.info(
        """
        Next upgrade:
        Automatic map insertion,
        north arrow,
        scale bar,
        PNG export,
        PDF export.
        """
)

# ==============================
# GENERATE LAYOUT
# ==============================

    if st.button(
        "Generate Layout",
        key="generate_layout"
):

        st.success(
            "Layout generated successfully!"
    )

    st.markdown(
        f"""
        # {map_title}

        {subtitle}

        -----------------------------

        🗺️ MAP LAYOUT PREVIEW

        Map Type:
        {map_type}

        Paper Size:
        {paper_size}

        -----------------------------
        """
    )

    if show_legend:
        st.write("✅ Legend Included")

    if show_north_arrow:
        st.write("✅ North Arrow Included")

    if show_scale_bar:
        st.write("✅ Scale Bar Included")


# ======================================
# REPORT ANALYSIS PAGE
# ======================================

if page == "Report Analysis":

    st.title(
        "Urban Climate Report Generator"
    )

    # ==============================
    # TITLE
    # ==============================

    doc.add_heading(
        'Urban Climate Intelligence Report',
        level=1
    )

    # ==============================
    # STUDY OVERVIEW
    # ==============================

    doc.add_heading(
        'Study Overview',
        level=2
    )

    doc.add_paragraph(
        f"Study Area: {study_area}"
    )

    doc.add_paragraph(
        f"Analysis Period: {report_period}"
    )

    # ==============================
    # DYNAMIC LST INTERPRETATION
    # ==============================

    if mean_lst >= 40:

        lst_interpretation = f"""

        The study area recorded a very high average
        land surface temperature of {mean_lst} °C,
        indicating severe urban heat conditions.

        Several parts of the area are likely experiencing
        excessive heat buildup due to dense buildings,
        roads, concrete surfaces, and limited vegetation.

        These conditions may increase heat discomfort,
        environmental stress, and cooling demand,
        especially during hotter periods.
        """

    elif mean_lst >= 35:

        lst_interpretation = f"""

        The study area recorded an average land surface
        temperature of {mean_lst} °C, suggesting
        moderately high urban temperature conditions.

        Some developed areas appear warmer than
        surrounding environments, likely due to
        urban expansion and reduced vegetation cover.

        Environmental cooling interventions may help
        reduce future heat intensity.
        """

    else:

        lst_interpretation = f"""

        The study area recorded a relatively moderate
        average temperature of {mean_lst} °C.

        Temperature conditions appear more stable,
        suggesting lower urban heat intensity and
        better environmental balance compared to
        heavily urbanized environments.
        """

    doc.add_heading(
        'Land Surface Temperature Analysis',
        level=2
    )

    doc.add_paragraph(
        lst_interpretation
    )

    # ==============================
    # DYNAMIC NDVI INTERPRETATION
    # ==============================

    if mean_ndvi < 0.2:

        ndvi_interpretation = f"""

        The study area recorded a low average NDVI
        value of {mean_ndvi}, indicating limited
        vegetation cover across several locations.

        Reduced vegetation may be contributing
        significantly to increasing urban heat
        conditions and environmental stress.
        """

    elif mean_ndvi < 0.5:

        ndvi_interpretation = f"""

        The study area recorded a moderate NDVI
        value of {mean_ndvi}, suggesting varying
        vegetation conditions across the environment.

        Some areas maintain healthy vegetation,
        while others show signs of reduced green cover.
        """

    else:

        ndvi_interpretation = f"""

        The study area recorded a relatively high
        NDVI value of {mean_ndvi}, indicating
        healthy vegetation presence across many areas.

        Existing vegetation appears to contribute
        positively to environmental cooling and
        ecological stability.
        """

    doc.add_heading(
        'Vegetation Analysis',
        level=2
    )

    doc.add_paragraph(
        ndvi_interpretation
    )

    # ==============================
    # DYNAMIC CORRELATION ANALYSIS
    # ==============================

    if correlation <= -0.7:

        correlation_interpretation = f"""

        The analysis recorded a strong negative
        correlation value of {correlation}
        between vegetation cover and land surface
        temperature.

        This indicates that areas with higher
        vegetation cover generally experienced
        significantly lower temperatures.
        """

    elif correlation <= -0.3:

        correlation_interpretation = f"""

        The analysis recorded a moderate negative
        correlation value of {correlation}
        between vegetation and surface temperature.

        Vegetation contributes to environmental
        cooling in several parts of the study area.
        """

    else:

        correlation_interpretation = f"""

        The analysis recorded a weak relationship
        value of {correlation} between vegetation
        and temperature conditions.

        Other environmental and urban factors may
        also contribute significantly to heat patterns.
        """

    doc.add_heading(
        'Correlation Analysis',
        level=2
    )

    doc.add_paragraph(
        correlation_interpretation
    )

    # ==============================
    # DYNAMIC IMPLICATIONS
    # ==============================

    if mean_lst >= 40 and mean_ndvi < 0.2:

        implication_text = """

        The environmental assessment indicates
        severe urban heat conditions combined
        with limited vegetation cover.

        These conditions may increase heat
        discomfort, environmental stress,
        and climate vulnerability.
        """

    elif mean_lst >= 35:

        implication_text = """

        The findings suggest moderately elevated
        urban temperature conditions across
        several parts of the study area.

        Environmental pressure may increase if
        sustainable interventions are not introduced.
        """

    else:

        implication_text = """

        Environmental conditions appear relatively
        stable compared to highly urbanized
        environments experiencing severe heat stress.
        """

    doc.add_heading(
        'Environmental Implications',
        level=2
    )

    doc.add_paragraph(
        implication_text
    )

    # ==============================
    # DYNAMIC RECOMMENDATIONS
    # ==============================

    if mean_lst >= 40 and mean_ndvi < 0.2:

        recommendation_text = """

        • Increase urban tree planting initiatives.

        • Protect and restore green spaces.

        • Reduce excessive concrete expansion.

        • Promote climate-sensitive urban planning.

        • Develop urban cooling strategies.

        • Strengthen environmental sustainability
          programs and awareness campaigns.
        """

    elif mean_lst >= 35:

        recommendation_text = """

        • Expand urban greenery and parks.

        • Strengthen sustainable land-use planning.

        • Protect existing vegetation cover.

        • Improve environmental monitoring systems.
        """

    else:

        recommendation_text = """

        • Maintain existing vegetation cover.

        • Continue sustainable urban development.

        • Promote long-term climate resilience planning.
        """

    doc.add_heading(
        'Recommendations',
        level=2
    )

    doc.add_paragraph(
        recommendation_text
    )

    # ==============================
    # CONCLUSION
    # ==============================

    doc.add_heading(
        'Conclusion',
        level=2
    )

    doc.add_paragraph(

        """
        The assessment highlights the growing
        importance of integrating environmental
        sustainability into urban planning and
        development processes.

        The results demonstrate that vegetation
        plays an important role in reducing
        urban heat conditions and improving
        climate resilience.
        """
    )

    # ==============================
    # SAVE DOCUMENT
    # ==============================

    doc.save(
        report_path
    )

    # ==============================
    # DOWNLOAD BUTTON
    # ==============================

    with open(
        report_path,
        "rb"
    ) as doc_file:

        st.download_button(
            label="Download Word Report",
            data=doc_file,
            file_name="Urban_Climate_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    st.success(
        "Word report generated successfully!"
    )

