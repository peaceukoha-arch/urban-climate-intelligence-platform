
# ======================================
# IMPORT LIBRARIES
# ======================================

import zipfile
import os
import tempfile
import ee
import geopandas as gpd
import pandas as pd
import streamlit as st
import matplotlib.pyplot as plt
import folium
import numpy as np


from streamlit_folium import st_folium
from datetime import date
from docx import Document 
from docx.shared import Inches


# ======================================
# INITIALIZE EARTH ENGINE
# ======================================

try:

    ee.Initialize()

except Exception:

    st.error(
        "Earth Engine authentication is not configured."
    )

    st.stop()


# ======================================
# LANDSAT 8 + 9 COLLECTION
# ======================================

landsat8 = ee.ImageCollection(
    'LANDSAT/LC08/C02/T1_L2'
)

landsat9 = ee.ImageCollection(
    'LANDSAT/LC09/C02/T1_L2'
)

landsat = landsat8.merge(
    landsat9
)


# ======================================
# PAGE CONFIG
# ======================================

st.set_page_config(
    page_title="Urban Heat Intelligence",
    layout="wide"
)

st.markdown("""
<style>

[data-testid="stAppViewContainer"] {
    background-color: #dcfce7  !important;
}

[data-testid="stMetric"] {
    background-color: green;
    border-radius: 15px;
    padding: 15px;
    box-shadow: 0px 2px 8px rgba(0,0,0,0.1);
}

</style>
""", unsafe_allow_html=True)

# ======================================
# CUSTOM STYLING
# ======================================

st.markdown("""
<style>

/* Sidebar background */

[data-testid="stSidebar"] {
    background: linear-gradient(
        180deg,
        #14532d,
        #166534,
        #15803d
    );
}

/* Labels */

[data-testid="stSidebar"] label {
    color: white !important;
    font-weight: 600;
}

/* Headings */

[data-testid="stSidebar"] h1,
[data-testid="stSidebar"] h2,
[data-testid="stSidebar"] h3,
[data-testid="stSidebar"] p {
    color: white !important;
}

/* Text boxes */

[data-testid="stSidebar"] input {
    color: black !important;
    background-color: white !important;
}

/* Select boxes */

[data-testid="stSidebar"] select {
    color: black !important;
    background-color: white !important;
}

</style>
""", unsafe_allow_html=True)

st.markdown("""
<style>

/* Upload button */

[data-testid="stFileUploader"] section {
    background-color: white !important;
    border: 2px solid #000000 !important;
    border-radius: 10px;
}

/* Browse files button */

[data-testid="stFileUploader"] button {
    background-color: black !important;
    color: white !important;
    border-radius: 8px !important;
    border: none !important;
}

/* Hover effect */

[data-testid="stFileUploader"] button:hover {
    background-color: #333333 !important;
    color: white !important;
}

</style>
""", unsafe_allow_html=True)

# ======================================
# SIDEBAR
# ======================================

st.sidebar.title(
    "Urban Climate Controls"
)

st.sidebar.write(
    "Configure analysis settings"
)


# ======================================
# NAVIGATION
# ======================================


page = st.sidebar.radio(
    "Navigation",
    [
        "Map Analysis",
        "Time-Series Analytics",
        "Report Analysis"
    ]
)
# ======================================
# MAP ANALYSIS PAGE
# ======================================

if page == "Map Analysis":

    st.markdown("""
    <div style="
    background: linear-gradient(
    90deg,
    #166534,
    #22c55e,
    #4ade80
    );
    padding:25px;
    border-radius:20px;
    color:white;
    text-align:center;
    ">

    <h1>
    🌍 Urban Climate Intelligence Platform
    </h1>

    <p>
    Transforming Data into Decisions 
    </p>

    </div>
    """,
    unsafe_allow_html=True)

    st.write(
        """
        Upload GeoJSON, KML, or ZIP Shapefile
        to generate:
        - Land Surface Temperature
        - NDVI
        - Hotspot Analysis
        """
    )

    # ==================================
    # DATE RANGE
    # ==================================

    start_date = st.sidebar.date_input(
        "Start Date",
        date(2024, 1, 1)
    )

    end_date = st.sidebar.date_input(
        "End Date",
        date(2024, 12, 31)
    )

    # ==================================
    # BASEMAP
    # ==================================

    basemap = st.sidebar.selectbox(
        "Select Basemap",
        [
            "OpenStreetMap",
            "SATELLITE",
            "HYBRID",
            "TERRAIN"
        ]
    )

    # ==================================
    # LAYER TOGGLES
    # ==================================

    show_lst = st.sidebar.checkbox(
        "Show LST",
        value=True
    )
    
    show_ndvi = st.sidebar.checkbox(
        "Show NDVI",
        value=True
    )

    show_hotspot = st.sidebar.checkbox(
        "Show Hotspot",
        value=True
    )

    # ==================================
    # FILE UPLOAD
    # ==================================

    uploaded_file = st.sidebar.file_uploader(
        "Upload Boundary",
        type=["geojson", "kml", "zip"]
    )

    # ==================================
    # PROCESS FILE
    # ==================================

    if uploaded_file is not None:

        file_extension = os.path.splitext(
            uploaded_file.name
        )[1]

        temp_file = tempfile.NamedTemporaryFile(
            delete=False,
            suffix=file_extension
        )

        temp_file.write(
            uploaded_file.read()
        )

        temp_file.close()

        # ==============================
        # READ FILES
        # ==============================

        if uploaded_file.name.endswith(
            ".geojson"
        ):

            gdf = gpd.read_file(
                temp_file.name
            )

        elif uploaded_file.name.endswith(
            ".kml"
        ):

            gdf = gpd.read_file(
                temp_file.name,
                driver="KML"
            )

        elif uploaded_file.name.endswith(
            ".zip"
        ):

            extract_path = tempfile.mkdtemp()

            with zipfile.ZipFile(
                temp_file.name,
                "r"
            ) as zip_ref:

                zip_ref.extractall(
                    extract_path
                )

            shp_file = None

            for root, dirs, files in os.walk(
                extract_path
            ):

                for file in files:

                    if file.endswith(".shp"):

                        shp_file = os.path.join(
                            root,
                            file
                        )

            gdf = gpd.read_file(
                shp_file
            )

        st.success(
            "Boundary uploaded successfully!"
        )

        # ==============================
        # EARTH ENGINE GEOMETRY
        # ==============================

        geojson = gdf.__geo_interface__

        ee_geometry = ee.FeatureCollection(
            geojson
        )

        # ==============================
        # CLOUD MASK
        # ==============================

        def mask_clouds(image):

            qa = image.select(
                'QA_PIXEL'
            )

            cloud = qa.bitwiseAnd(
                1 << 3
            ).eq(0)

            cloud_shadow = qa.bitwiseAnd(
                1 << 4
            ).eq(0)

            mask = cloud.And(
                cloud_shadow
            )

            return image.updateMask(mask)

        # ==============================
        # CALCULATE INDICES
        # ==============================

        def calculate_indices(image):

            thermal = image.select(
                'ST_B10'
            ).multiply(
                0.00341802
            ).add(
                149.0
            )

            lst = thermal.subtract(
                273.15
            ).rename(
                'LST'
            )

            ndvi = image.normalizedDifference(
                ['SR_B5', 'SR_B4']
            ).rename(
                'NDVI'
            )

            return image.addBands([
                lst,
                ndvi
            ])

        # ==============================
        # PROCESS COLLECTION
        # ==============================

        processed = landsat \
            .filterBounds(
                ee_geometry
            ) \
            .filterDate(
                str(start_date),
                str(end_date)
            ) \
            .filter(
                ee.Filter.lt(
                    'CLOUD_COVER',
                    20
                )
            ) \
            .map(
                mask_clouds
            ) \
            .map(
                calculate_indices
            )

        # ==============================
        # CREATE COMPOSITES
        # ==============================

        lst_image = processed.select(
            'LST'
        ).median().clip(
            ee_geometry
        )

        lst_stats = lst_image.reduceRegion(
            reducer=ee.Reducer.minMax(),
            geometry=ee_geometry,
            scale=30,
            maxPixels=1e13
        )

        lst_min = lst_stats.getInfo()['LST_min']
        lst_max = lst_stats.getInfo()['LST_max']
        
        lst_mean_stats = lst_image.reduceRegion(
            reducer=ee.Reducer.mean(),
            geometry=ee_geometry,
            scale=30,
            maxPixels=1e13
        )

        mean_lst = lst_mean_stats.getInfo()['LST']

        st.session_state["mean_lst"] = mean_lst
        
        # Save for Report Generator
        st.session_state["lst_min"] = lst_min
        st.session_state["lst_max"] = lst_max

        interval = (lst_max - lst_min) / 4

        class1 = lst_min + interval
        class2 = lst_min + (interval * 2)
        class3 = lst_min + (interval * 3)


        ndvi_image = processed.select(
            'NDVI'
        ).median().clip(
            ee_geometry
        )

        ndvi_stats = ndvi_image.reduceRegion(
            reducer=ee.Reducer.mean(),
            geometry=ee_geometry,
            scale=30,
            maxPixels=1e13
        )

        mean_ndvi = ndvi_stats.getInfo()['NDVI']

        st.session_state["mean_ndvi"] = mean_ndvi
        
        # ==============================
        # HOTSPOT CLASSIFICATION
        # ==============================

        hotspot = lst_image.expression(

            "(b('LST') < 25) ? 1" +
            ": (b('LST') < 30) ? 2" +
            ": (b('LST') < 35) ? 3" +
            ": (b('LST') < 40) ? 4" +
            ": 5"

        ).rename(
            'Hotspot'
        ).clip(
            ee_geometry
        )

        # ==============================
        # LST LEGEND COLORS
        # ==============================

        st.sidebar.subheader(
            "LST Legend Colors"
        )

        lst_low = st.sidebar.color_picker(
            "Low",
            "#00FF00"
        )
        
        lst_moderate = st.sidebar.color_picker(
            "Moderate",
            "#FFFF00"
        )
    
        lst_high = st.sidebar.color_picker(
            "High",
            "#FFA500"
        )

        lst_extreme = st.sidebar.color_picker(
            "Extreme",
            "#FF0000"
        )

        # ==============================
        # NDVI LEGEND COLORS
        # ==============================

        st.sidebar.subheader(
            "NDVI Legend Colors"
        )

        ndvi_low = st.sidebar.color_picker(
            "Low Vegetation",
            "#8B4513"
        )

        ndvi_moderate = st.sidebar.color_picker(
            "Moderate Vegetation",
            "#FFFF00"
        )

        ndvi_high = st.sidebar.color_picker(
            "High Vegetation",
            "#008000"
        )

        # ==============================
        # HOTSPOT LEGEND COLORS
        # ==============================

        st.sidebar.subheader(
            "Hotspot Legend Colors"
        )

        hotspot_low = st.sidebar.color_picker(
            "Low Hotspot",
            "#abd9e9"
        )

        hotspot_moderate = st.sidebar.color_picker(
            "Moderate Hotspot",
            "#ffffbf"
        )

        hotspot_high = st.sidebar.color_picker(
            "High Hotspot",
            "#fdae61"
        )

        hotspot_extreme = st.sidebar.color_picker(
            "Extreme Hotspot",
            "#d7191c"
        )

        # ==============================
        # VISUALIZATION PARAMETERS
        # ==============================

        lst_vis = {
            'min': 1,
            'max': 4,
            'palette': [
                str(lst_low).replace("#", ""),
                str(lst_moderate).replace("#", ""),
                str(lst_high).replace("#", ""),
                str(lst_extreme).replace("#", "")
            ]
        }

        ndvi_vis = {
            'min': -1,
            'max': 1,
            'palette': [
                str(ndvi_low).replace("#", ""),
                str(ndvi_moderate).replace("#", ""),
                str(ndvi_high).replace("#", "")
            ]
        }

        hotspot_vis = {
            'min': 1,
            'max': 4,
            'palette': [
                str(hotspot_low).replace("#", ""),
                str(hotspot_moderate).replace("#", ""),
                str(hotspot_high).replace("#", ""),
                str(hotspot_extreme).replace("#", "")
            ]
        }

        # ==============================
        # MAP CENTER
        # ==============================

        centroid = gdf.geometry.centroid.iloc[0]

        map_center = [
        centroid.y,
        centroid.x
        ]
        # ==============================
        # CREATE MAP
        # ==============================

        # ======================================
        # CREATE MAP
        # ======================================

        if basemap == "OpenStreetMap":

            m = folium.Map(
                location=map_center,
                zoom_start=8,
                control_scale=True,
                tiles="OpenStreetMap"
            )

        elif basemap == "SATELLITE":

            m = folium.Map(
                location=map_center,
                zoom_start=8,
                control_scale=True,
                tiles=None
            )

            folium.TileLayer(
                tiles="https://mt1.google.com/vt/lyrs=s&x={x}&y={y}&z={z}",
                attr="Google",
                name="Satellite"
            ).add_to(m)

        elif basemap == "HYBRID":

            m = folium.Map(
                location=map_center,
                zoom_start=8,
                control_scale=True,
                tiles=None
            )

            folium.TileLayer(
                tiles="https://mt1.google.com/vt/lyrs=y&x={x}&y={y}&z={z}",
                attr="Google",
                name="Hybrid"
            ).add_to(m)

        elif basemap == "TERRAIN":

            m = folium.Map(
                location=map_center,
                zoom_start=8,
                control_scale=True,
                tiles=None
            )

            folium.TileLayer(
                tiles="https://mt1.google.com/vt/lyrs=p&x={x}&y={y}&z={z}",
                attr="Google",
                name="Terrain"
            ).add_to(m)

        # ==============================
        # BASEMAPS
        # ==============================

        if basemap == "OpenStreetMap":

            folium.TileLayer(
                'OpenStreetMap',
                name='OpenStreetMap'
            ).add_to(m)

        elif basemap == "SATELLITE":

            folium.TileLayer(
                tiles='https://mt1.google.com/vt/lyrs=s&x={x}&y={y}&z={z}',
                attr='Google',
                name='Satellite'
            ).add_to(m)

        elif basemap == "HYBRID":

            folium.TileLayer(
                tiles='https://mt1.google.com/vt/lyrs=y&x={x}&y={y}&z={z}',
                attr='Google',
                name='Hybrid'
            ).add_to(m)

        elif basemap == "TERRAIN":

            folium.TileLayer(
                tiles='https://mt1.google.com/vt/lyrs=p&x={x}&y={y}&z={z}',
                attr='Google',
                name='Terrain'
            ).add_to(m)

        # ==============================
        # STUDY AREA OUTLINE
        # ==============================

        folium.GeoJson(
            geojson,
            name="Study Area",
            style_function=lambda x: {
                'fillColor': 'none',
                'color': 'black',
                'weight': 2
            }
        ).add_to(m)

        # ==============================
        # CLASSIFIED LST
        # ==============================

        classified_lst = (
            lst_image.gt(lst_min).And(
                lst_image.lte(class1)
            ).multiply(1)
            .add(
                lst_image.gt(class1).And(
                    lst_image.lte(class2)
                ).multiply(2)
            )
            .add(
                lst_image.gt(class2).And(
                    lst_image.lte(class3)
                ).multiply(3)
            )
            .add(
                lst_image.gt(class3).multiply(4)
            )
        )

        
        # ==============================
        # LST LAYER
        # ==============================

        if show_lst:

            map_id = classified_lst.getMapId(
                lst_vis
        )

            folium.TileLayer(
                tiles=map_id['tile_fetcher'].url_format,
                attr='Google Earth Engine',
                name='LST',
                overlay=True,
                control=True
            ).add_to(m)

        # ==============================
        # NDVI LAYER
        # ==============================

        if show_ndvi:

            map_id = ee.Image(
                ndvi_image
            ).getMapId(ndvi_vis)

            folium.TileLayer(
                tiles=map_id['tile_fetcher'].url_format,
                attr='Google Earth Engine',
                name='NDVI',
                overlay=True,
                control=True
            ).add_to(m)

        # ==============================
        # HOTSPOT LAYER
        # ==============================

        if show_hotspot:

            map_id = ee.Image(
                hotspot
            ).getMapId(
                hotspot_vis
            )

            folium.TileLayer(
                tiles=map_id['tile_fetcher'].url_format,
                attr='Google Earth Engine',
                name='Hotspot',
                overlay=True,
                control=True
            ).add_to(m)

        # ==============================
        # LAYER CONTROL
        # ==============================

        folium.LayerControl().add_to(m)

        # ==============================
        # COMBINED LEGENDS
        # ==============================

        legend_html = ""

        # ==============================
        # LST LEGEND
        # ==============================

        
        if show_lst:

            legend_html += f"""

            <div style="
            position: fixed;
            bottom: 50px;
            right: 100px;
            width: 180px;
            height: 150px;
            background-color: white;
            border:2px solid grey;
            z-index:9999;
            font-size:14px;
            padding: 10px;
            border-radius:8px;
            ">

            <b>LST (°C)</b><br><br>

            <i style="background:{lst_low};
            width:15px;
            height:15px;
            float:left;
            margin-right:8px;"></i>
            Low<br>

            <i style="background:{lst_moderate};
            width:15px;
            height:15px;
            float:left;
            margin-right:8px;"></i>
            Moderate<br>

            <i style="background:{lst_high};
            width:15px;
            height:15px;
            float:left;
            margin-right:8px;"></i>
            High<br>

            <i style="background:{lst_extreme};
            width:15px;
            height:15px;
            float:left;
            margin-right:8px;"></i>
            Extreme

            </div>

            """


        # ==============================
        # NDVI LEGEND
        # ==============================

        if show_ndvi:

            legend_html += f"""

            <div style="
            position: fixed;
            bottom: 50px;
            right: 100px;
            width: 180px;
            height: 130px;
            background-color: white;
            border:2px solid grey;
            z-index:9999;
            font-size:14px;
            padding: 10px;
            border-radius:8px;
            ">

            <b>NDVI</b><br><br>

            <i style="background:{ndvi_low};
            width:15px;
            height:15px;
            float:left;
            margin-right:8px;"></i>
            Low Vegetation<br>

            <i style="background:{ndvi_moderate};
            width:15px;
            height:15px;
            float:left;
            margin-right:8px;"></i>
            Moderate<br>

            <i style="background:{ndvi_high};
            width:15px;
            height:15px;
            float:left;
            margin-right:8px;"></i>
            High Vegetation

            </div>

            """

        # ==============================
        # HOTSPOT LEGEND
        # ==============================

        if show_hotspot:

            legend_html += f"""

            <div style="
            position: fixed;
            bottom: 50px;
            right: 100px;
            width: 180px;
            height: 170px;
            background-color: white;
            border:2px solid grey;
            z-index:9999;
            font-size:14px;
            padding: 10px;
            border-radius:8px;
            ">

            <b>Heat Hotspots</b><br><br>

            <i style="background:{hotspot_low};
            width:15px;
            height:15px;
            float:left;
            margin-right:8px;"></i>
            Low<br>

            <i style="background:{hotspot_moderate};
            width:15px;
            height:15px;
            float:left;
            margin-right:8px;"></i>
            Moderate<br>

            <i style="background:{hotspot_high};
            width:15px;
            height:15px;
            float:left;
            margin-right:8px;"></i>
            High<br>

            <i style="background:{hotspot_extreme};
            width:15px;
            height:15px;
            float:left;
            margin-right:8px;"></i>
            Extreme

            </div>

            """

        m.get_root().html.add_child(
            folium.Element(
                legend_html
            )
        )

        # ==============================
        # DISPLAY MAP
        # ==============================

        if uploaded_file is not None:

            st_folium(
                m,
                width=1200,
                height=700,
                key=(
                    f"{lst_low}_{lst_moderate}_{lst_high}_{lst_extreme}_"
                    f"{ndvi_low}_{ndvi_moderate}_{ndvi_high}_"
                    f"{hotspot_low}_{hotspot_moderate}_{hotspot_high}_{hotspot_extreme}"
                )
            )


        st.markdown("---")

        st.subheader(
            "Analysis Summary"
        )

        col1, col2, col3 = st.columns(3)

        # ==============================
        # LST SUMMARY
        # ==============================

        if show_lst:

            with col1:

                st.metric(
                    "Minimum LST (°C)",
                    f"{lst_min:.2f}"
                )

                st.metric(
                    "Maximum LST (°C)",
                    f"{lst_max:.2f}"
                )

# ==============================
# NDVI SUMMARY
# ==============================

        if show_ndvi:

            ndvi_stats = ndvi_image.reduceRegion(
                reducer=ee.Reducer.mean(),
                geometry=ee_geometry,
                scale=30,
                maxPixels=1e13
            )

            ndvi_mean = ndvi_stats.getInfo().get(
                "NDVI",
                0
            )

            with col2:

                st.metric(
                    "Mean NDVI",
                    f"{ndvi_mean:.2f}"
            )

# ==============================
# HOTSPOT SUMMARY
# ==============================

        if show_hotspot:

            hotspot_stats = hotspot.reduceRegion(
                reducer=ee.Reducer.mean(),
                geometry=ee_geometry,
                scale=30,
                maxPixels=1e13
            )

            hotspot_mean = hotspot_stats.getInfo().get(
                "Hotspot",
                0
            )

            st.button(
                "Download Map PNG",
                key="download_map_png"
            )

        
# ======================================
# TIMESERIES PAGE
# ======================================

if page == "Time-Series Analytics":

    st.title(
        "Climate Time-Series Analytics"
    )

    st.write(
        """
        Analyze:
        - Yearly Trends
        - Monthly Trends
        - Correlation Analysis
        """
    )

    # ==================================
    # TIMESERIES TYPE
    # ==================================

    timeseries_type = st.sidebar.radio(
        "Time-Series Type",
        [
            "Yearly",
            "Monthly"
        ]
    )

    # ==================================
    # ANALYSIS TYPE
    # ==================================

    analysis_type = st.sidebar.selectbox(
        "Analysis Type",
        [
            "LST Trend",
            "NDVI Trend",
            "LST-NDVI Correlation"
        ]
    )

    # ==================================
    # YEARLY SETTINGS
    # ==================================

    if timeseries_type == "Yearly":

        start_year = st.sidebar.selectbox(
            "Start Year",
            list(range(2015, 2026)),
            index=0
        )

        end_year = st.sidebar.selectbox(
            "End Year",
            list(range(2015, 2026)),
            index=10
        )

    # ==================================
    # MONTHLY SETTINGS
    # ==================================

    elif timeseries_type == "Monthly":

        start_month = st.sidebar.date_input(
            "Start Month",
            date(2015, 1, 1)
        )

        end_month = st.sidebar.date_input(
            "End Month",
            date.today()
        )

    # ==================================
    # FILE UPLOAD
    # ==================================

    uploaded_file = st.sidebar.file_uploader(
        "Upload Boundary",
        type=["geojson", "kml", "zip"],
        key="timeseries_upload"
    )

    # ==================================
    # PROCESS FILE
    # ==================================

    if uploaded_file is not None:

        file_extension = os.path.splitext(
            uploaded_file.name
        )[1]

        temp_file = tempfile.NamedTemporaryFile(
            delete=False,
            suffix=file_extension
        )

        temp_file.write(
            uploaded_file.read()
        )

        temp_file.close()

        # ==============================
        # READ FILES
        # ==============================

        if uploaded_file.name.endswith(".geojson"):

            gdf = gpd.read_file(
                temp_file.name
            )

        elif uploaded_file.name.endswith(".kml"):

            gdf = gpd.read_file(
                temp_file.name,
                driver="KML"
            )

        elif uploaded_file.name.endswith(".zip"):

            extract_path = tempfile.mkdtemp()

            with zipfile.ZipFile(
                temp_file.name,
                "r"
            ) as zip_ref:

                zip_ref.extractall(
                    extract_path
                )

            shp_file = None

            for root, dirs, files in os.walk(
                extract_path
            ):

                for file in files:

                    if file.endswith(".shp"):

                        shp_file = os.path.join(
                            root,
                            file
                        )

            gdf = gpd.read_file(
                shp_file
            )

        st.success(
            "Boundary uploaded successfully!"
        )

        # ======================================
        # SIMPLIFY LARGE BOUNDARIES
        # ======================================

        gdf = gdf.to_crs(3857)

        gdf["geometry"] = gdf.geometry.simplify(
            tolerance=5000
        )

        gdf = gdf.to_crs(4326)

        # ==============================
        # EARTH ENGINE GEOMETRY
        # ==============================

        geojson = gdf.__geo_interface__

        ee_geometry = ee.FeatureCollection(
            geojson
        )

        # ==============================
        # CLOUD MASK
        # ==============================

        def mask_clouds(image):

            qa = image.select(
                'QA_PIXEL'
            )

            cloud = qa.bitwiseAnd(
                1 << 3
            ).eq(0)

            cloud_shadow = qa.bitwiseAnd(
                1 << 4
            ).eq(0)

            mask = cloud.And(
                cloud_shadow
            )

            return image.updateMask(mask)

        # ==============================
        # CALCULATE INDICES
        # ==============================

        def calculate_indices(image):

            thermal = image.select(
                'ST_B10'
            ).multiply(
                0.00341802
            ).add(
                149.0
            )

            lst = thermal.subtract(
                273.15
            ).rename(
                'LST'
            )

            ndvi = image.normalizedDifference(
                ['SR_B5', 'SR_B4']
            ).rename(
                'NDVI'
            )

            return image.addBands([
                lst,
                ndvi
            ])

        # ==============================
        # PERIODS
        # ==============================

        if timeseries_type == "Yearly":

            periods = list(
                range(start_year, end_year + 1)
            )

            x_values = periods

        else:

            monthly_dates = pd.date_range(
                start=start_month,
                end=end_month,
                freq='MS'
            )

            periods = monthly_dates

            x_values = [
                d.strftime("%b-%Y")
                for d in monthly_dates
            ]

        lst_values = []

        ndvi_values = []

        # ==============================
        # ANALYSIS LOOP
        # ==============================

        with st.spinner(
            "Generating analysis..."
        ):

            for period in periods:

                if timeseries_type == "Yearly":

                    start = f"{period}-01-01"

                    end = f"{period}-12-31"

                else:

                    start = period.strftime(
                        "%Y-%m-%d"
                    )

                    if period.month == 12:

                        next_month = period.replace(
                            year=period.year + 1,
                            month=1
                        )

                    else:

                        next_month = period.replace(
                            month=period.month + 1
                        )

                    end = next_month.strftime(
                        "%Y-%m-%d"
                    )

                collection = landsat \
                    .filterBounds(
                        ee_geometry
                    ) \
                    .filterDate(
                        start,
                        end
                    ) \
                    .filter(
                        ee.Filter.lt(
                            'CLOUD_COVER',
                            20
                        )
                    ) \
                    .map(
                        mask_clouds
                    ) \
                    .map(
                        calculate_indices
                    )

                # ==========================
                # LST
                # ==========================

                lst_image = collection.select(
                    'LST'
                ).median()

                lst_stats = lst_image.reduceRegion(
                    reducer=ee.Reducer.mean(),
                    geometry=ee_geometry,
                    scale=500,
                    bestEffort=True
                )

                lst_mean = lst_stats.getInfo().get(
                    'LST',
                    None
                )

                if lst_mean is None:

                    if len(lst_values) > 0:

                        lst_mean = round(
                            sum(lst_values) / len(lst_values),
                            2
                        )

                    else:

                        lst_mean = 0

                else:

                    lst_mean = round(
                        lst_mean,
                        2
                    )

                lst_values.append(
                    lst_mean
                )

                # ==========================
                # NDVI
                # ==========================

                ndvi_image = collection.select(
                    'NDVI'
                ).median()

                ndvi_stats = ndvi_image.reduceRegion(
                    reducer=ee.Reducer.mean(),
                    geometry=ee_geometry,
                    scale=500,
                    bestEffort=True
                )

                ndvi_mean = ndvi_stats.getInfo().get(
                    'NDVI',
                    None
                )

                if ndvi_mean is None:

                    if len(ndvi_values) > 0:

                        ndvi_mean = round(
                            sum(ndvi_values) / len(ndvi_values),
                            3
                        )

                    else:

                        ndvi_mean = 0

                else:

                    ndvi_mean = round(
                        ndvi_mean,
                        3
                    )

                ndvi_values.append(
                    ndvi_mean
                )

        # ==============================
        # DATAFRAME
        # ==============================
        if analysis_type == "LST Trend":

            df = pd.DataFrame({

                'Period': x_values,

                'Mean_LST': lst_values,
            })

        elif analysis_type == "NDVI Trend":

            df = pd.DataFrame({
                'Period': x_values,
                'Mean_NDVI': ndvi_values
            })

        else:  # LST-NDVI Correlation

            df = pd.DataFrame({
                'Mean_LST': lst_values,
                'Mean_NDVI': ndvi_values
            })

        # ==============================
        # SAVE FOR REPORT GENERATOR
        # ==============================

        st.session_state["timeseries_periods"] = x_values

        st.session_state["timeseries_lst"] = lst_values

        st.session_state["timeseries_ndvi"] = ndvi_values

        st.subheader(
            "Time-Series Data"
        )

        st.dataframe(df)

        # ==============================
        # PLOT
        # ==============================

        if len(x_values) <= 15:

            fig, ax = plt.subplots(
                figsize=(12, 6)
            )

        elif len(x_values) <= 60:

            fig, ax = plt.subplots(
                figsize=(20, 8)
            )

        else:

            fig, ax = plt.subplots(
                figsize=(30, 10)
            )

        if analysis_type == "LST Trend":

            ax.plot(
                x_values,
                lst_values,
                marker='o'
            )

            ax.set_ylabel(
                "Mean LST (°C)"
            )
            ax.set_title(
                "Land Surface Temperature Trend",
                fontsize=20,
                fontweight='bold'
            )

        elif analysis_type == "NDVI Trend":

            ax.plot(
                x_values,
                ndvi_values,
                marker='s'
            )

            ax.set_ylabel(
                "Mean NDVI"
            )
            
            ax.set_title(
                "Vegetation Trend",
                fontsize=20,
                fontweight='bold'
            )

        elif analysis_type == "LST-NDVI Correlation":

            correlation = df[
                'Mean_LST'
            ].corr(
                df['Mean_NDVI']
            )

            r2 = correlation ** 2

            col1, col2 = st.columns(2)

            with col1:

                st.metric(
                    "Correlation",
                    round(correlation, 3)
                 )

            with col2:

                st.metric(
                    "R²",
                    round(r2, 3)
                )

            ax.scatter(
                ndvi_values,
                lst_values
            )

            ax.set_xlabel(
                "Mean NDVI"
            )

            ax.set_ylabel(
                "Mean LST (°C)"
            )

            ax.set_title(
                "LST-NDVI Correlation",
                fontsize=20,
                fontweight='bold'
            )

            
        # ==============================
        # CORRELATION INTERPRETATION
        # ==============================

            if correlation <= -0.7:

                interpretation = """
                Strong negative relationship between vegetation
                cover and land surface temperature.

                Implication:
                Vegetation plays a major role in regulating
                surface temperature. Increasing green spaces
                and tree cover is likely to significantly
                reduce urban heat conditions.
                """

            elif correlation <= -0.4:

                interpretation = """
                Moderate negative relationship between vegetation
                cover and land surface temperature.

                Implication:
                Areas with higher vegetation cover generally
                experience lower temperatures. Expanding urban
                greenery could help reduce heat stress, although
                other factors such as buildings, roads, and bare
                surfaces also influence temperature patterns.
                """

            elif correlation <= -0.2:

                interpretation = """
                Weak negative relationship between vegetation
                cover and land surface temperature.

                Implication:
                Vegetation provides some cooling benefits, but
                its influence on temperature is limited. Other
                environmental and urban factors are likely
                contributing more strongly to temperature
                variation within the study area.
                """

            else:

                interpretation = """
                Little or no relationship detected between
                vegetation cover and land surface temperature.

                Implication:
                Changes in vegetation cover do not appear to
                strongly influence temperature patterns within
                the study area. Other factors such as urban
                development, impervious surfaces, industrial
                activities, topography, or climate conditions
                may be the primary drivers of temperature
                variation.
                """

            st.info(
                interpretation
            )
        # ==============================
        # FONT SETTINGS
            # ==============================

        ax.tick_params(
            axis='x',
            labelsize=14
        )

        ax.tick_params(
            axis='y',
            labelsize=14
        )

        plt.xticks(
            rotation=45,
            ha="right",
            fontsize=14
        )
        

# ==============================
# IMPROVED AXIS DISPLAY
# ==============================
        if len(x_values) > 24:

            step = max(
                1,
                len(x_values) // 20
            )

            ax.set_xticks(
                range(0, len(x_values), step)
            )

            ax.set_xticklabels(
                x_values[::step],
                rotation=45,
                ha="right"
            )

        else:

            plt.xticks(
                rotation=45,
                ha="right",
                fontsize=14
            )
        ax.grid(True)

        plt.tight_layout()

        st.pyplot(fig)

        # ==============================
        # EXPORT PNG
        # ==============================

        chart_path = "timeseries_chart.png"

        fig.savefig(
            chart_path,
            dpi=600,
            bbox_inches='tight'
        )

        with open(
            chart_path,
            "rb"
        ) as file:

            st.download_button(
                label="Download Chart PNG",
                data=file,
                file_name="Climate_Timeseries_Chart.png",
                mime="image/png"
            )

        # ==============================
        # EXPORT EXCEL
        # ==============================

        excel_path = "timeseries_data.xlsx"

        df.to_excel(
            excel_path,
            index=False
        )

        with open(
            excel_path,
            "rb"
        ) as file:

            st.download_button(
                label="Download Excel Data",
                data=file,
                file_name="Climate_Timeseries_Data.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )

    # =# ======================================
# REPORT ANALYSIS PAGE
# ======================================

if page == "Report Analysis":

    st.title(
        "Urban Climate Report Generator"
    )

    st.write(
        """
        Generate an automated urban climate
        intelligence report from analysis results.
        """
    )

    # ==============================
    # RETRIEVE ANALYSIS RESULTS
    # ==============================

    lst_min = st.session_state.get(
        "lst_min",
        None
    )

    lst_max = st.session_state.get(
        "lst_max",
        None
    )

    mean_lst = st.session_state.get(
        "mean_lst",
        None
    )

    mean_ndvi = st.session_state.get(
        "mean_ndvi",
        None
    )
    
    correlation = st.session_state.get(
        "correlation",
        None
    )
    
    timeseries_periods = st.session_state.get(
        "timeseries_periods",
        []
    )

    timeseries_lst = st.session_state.get(
        "timeseries_lst",
        []
    )

    timeseries_ndvi = st.session_state.get(
        "timeseries_ndvi",
        []
    )

    # ==============================
    # CHECK RESULTS EXIST
    # ==============================

    if mean_lst is None:

        st.warning(
            "Please run Map Analysis first."
        )

        st.stop()

    # ==============================
    # DISPLAY ANALYSIS SUMMARY
    # ==============================

    st.subheader(
        "Analysis Summary"
    )

    col1, col2, col3 = st.columns(3)

    with col1:

        st.metric(
            "Minimum LST (°C)",
            f"{lst_min:.2f}"
        )

        st.metric(
            "Maximum LST (°C)",
            f"{lst_max:.2f}"
        )

    with col2:

        st.metric(
            "Mean LST (°C)",
            f"{mean_lst:.2f}"
        )

    with col3:

        st.metric(
            "Mean NDVI",
            f"{mean_ndvi:.2f}"
        )
    # ==================================
    # PROJECT INFORMATION
    # ==================================

    study_area = st.text_input(
        "Study Area Name",
    )

    report_period = st.text_input(
        "Analysis Period",
    )

    st.subheader(
        "Project Information"
    )

    project_purpose = st.text_area(
        "Purpose of Assessment(Optional)"
    )
    
    
    # ==============================
    # GENERATE REPORT
    # ==============================

    if st.button(
        "Generate Word Report"
    ):

        report_path = "Urban_Climate_Report.docx"

        doc = Document()

        # ==============================
        # TITLE
        # ==============================

        doc.add_heading(
            "Urban Climate Intelligence Report",
            level=1
        )

        # ==============================
        # EXECUTIVE SUMMARY
        # ==============================

        doc.add_heading(
            "Executive Summary",
            level=2
        )

        doc.add_paragraph(
            f"""
            This report presents an assessment of land
            surface temperature and vegetation conditions
            within {study_area} during the period
            {report_period}.

            The analysis revealed temperatures ranging
            from {lst_min:.2f} °C to {lst_max:.2f} °C,
            with an average temperature of {mean_lst:.2f} °C.

            The mean NDVI value of {mean_ndvi:.2f}
            was used to evaluate vegetation conditions
            and their influence on urban thermal patterns.

            The findings provide insights to support
            climate resilience planning, environmental
            management, and sustainable urban development.
            """
        )
        # ==============================
        # STUDY OVERVIEW
        # ==============================

        doc.add_heading(
            'Study Overview',
            level=2
        )

        doc.add_paragraph(
            f"Study Area: {study_area}"
        )

        doc.add_paragraph(
            f"Analysis Period: {report_period}"
        )
        doc.add_paragraph(
            f"Purpose: {project_purpose}"
        )

        # ==============================
        # LST ANALYSIS
        # ==============================

        if mean_lst > 40:

            lst_interpretation = f"""
            The analysis revealed a minimum land surface
            temperature of {lst_min:.2f} °C and a maximum
            temperature of {lst_max:.2f} °C.

            The mean land surface temperature of
            {mean_lst:.2f} °C indicates severe urban heat
            conditions within the study area.

            These elevated temperatures suggest the
            presence of extensive heat hotspots and may
            increase environmental and public health risks.
            """

        elif mean_lst > 35:

            lst_interpretation = f"""
            The analysis revealed a minimum land surface
            temperature of {lst_min:.2f} °C and a maximum
            temperature of {lst_max:.2f} °C.

            The mean land surface temperature of
            {mean_lst:.2f} °C indicates moderate to high
            urban heat conditions across the study area.

            Temperature distribution suggests the presence
            of localized hotspots requiring climate-smart
            planning interventions.
            """

        else:

            lst_interpretation = f"""
            The analysis revealed a minimum land surface
            temperature of {lst_min:.2f} °C and a maximum
            temperature of {lst_max:.2f} °C.

            The mean land surface temperature of
            {mean_lst:.2f} °C indicates relatively moderate
            thermal conditions across the study area.

            Existing environmental conditions appear to
            provide some degree of temperature regulation.
            """

        doc.add_heading(
            "Land Surface Temperature Analysis",
            level=2
        )

        doc.add_paragraph(
            lst_interpretation
        )
        # ==============================
        # NDVI ANALYSIS
        # ==============================

        if mean_ndvi < 0.2:

            ndvi_interpretation = f"""
            The mean NDVI value of {mean_ndvi:.2f}
            indicates low vegetation cover across
            the study area.

            Limited vegetation may be contributing
            to increased land surface temperatures
            and the formation of urban heat hotspots.

            The observed vegetation condition suggests
            a need for enhanced urban greening and
            ecosystem restoration initiatives.
            """

        elif mean_ndvi < 0.5:

            ndvi_interpretation = f"""
            The mean NDVI value of {mean_ndvi:.2f}
            indicates moderate vegetation cover.

            Vegetation appears to provide some cooling
            benefits, although spatial variations may
            still contribute to localized heat stress.

            Strengthening green infrastructure could
            further improve environmental quality and
            climate resilience.
            """

        else:

            ndvi_interpretation = f"""
            The mean NDVI value of {mean_ndvi:.2f}
            indicates relatively high vegetation cover.

            Existing vegetation is likely contributing
            positively to temperature regulation,
            ecosystem stability, and environmental
            sustainability.

            Continued conservation efforts are
            recommended to maintain these benefits.
            """

        doc.add_heading(
            "Vegetation Analysis",
            level=2
        )

        doc.add_paragraph(
        ndvi_interpretation
        )

        # ==============================
        # TIME-SERIES ANALYSIS
        # ==============================

        doc.add_heading(
            "Time-Series Analysis",
            level=2
        )

        trend_text = """
        Time-series data is unavailable.
        Please run the Time-Series Analytics
        module before generating the report.
        """

        if len(timeseries_lst) > 1:

            first_lst = timeseries_lst[0]

            last_lst = timeseries_lst[-1]

            temperature_change = round(
            last_lst - first_lst,
            2
        )

            if temperature_change > 0:

                trend_text = f"""
                Time-series analysis revealed an
                increasing temperature trend.

                Mean land surface temperature
                increased from {first_lst:.2f} °C
                to {last_lst:.2f} °C,
                representing an increase of
                {temperature_change:.2f} °C.
                """

        elif temperature_change < 0:

            trend_text = f"""
            Time-series analysis revealed a
            decreasing temperature trend.

            Mean land surface temperature
            changed from {first_lst:.2f} °C
            to {last_lst:.2f} °C,
            representing a decrease of
            {abs(temperature_change):.2f} °C.
            """

        else:

            trend_text = """
            Time-series analysis indicates
            relatively stable temperature
            conditions throughout the
            study period.
            """

        doc.add_paragraph(
            trend_text
        )

        # ==============================
        # ENVIRONMENTAL IMPLICATIONS
        # ==============================

        if mean_lst > 40:

            implication = """
            The study area is experiencing severe urban heat
            conditions which may increase environmental stress,
            energy consumption, and heat-related risks.

            Such conditions may negatively affect public health,
            ecosystem stability, and urban livability.
            """

        elif mean_lst > 35:

            implication = """
            The study area exhibits moderate to high urban
            heat conditions requiring targeted mitigation
            measures and climate adaptation strategies.

            Continued urban expansion without adequate
            vegetation management may further intensify
            thermal stress.
            """

        else:

            implication = """
            The study area currently experiences relatively
            moderate thermal conditions with localized
            temperature variations.

            Existing environmental conditions appear to
            provide some degree of thermal regulation.
            """

        doc.add_heading(
            "Environmental Implications",
            level=2
        )

        doc.add_paragraph(
            implication
        )
        # ==============================
        # RECOMMENDATIONS
        # ==============================

        if mean_ndvi < 0.2:

            recommendation = """
            1. Increase urban tree planting and green
            infrastructure development.

            2. Protect and restore existing vegetation
            to improve environmental cooling.

            3. Establish urban green corridors and
            climate resilience programmes.

            4. Integrate nature-based solutions into
            urban development planning.
            """

        elif mean_ndvi < 0.5:

            recommendation = """
            1. Strengthen vegetation management and
            urban greening initiatives.

            2. Improve connectivity between existing
            green spaces.

            3. Promote sustainable land-use planning
            to minimize heat accumulation.

            4. Monitor vegetation health and thermal
            conditions regularly.
            """

        else:

            recommendation = """
            1. Maintain existing vegetation cover and
            ecosystem services.

            2. Encourage sustainable environmental
            management practices.

            3. Monitor future urban expansion to
            prevent vegetation loss.

            4. Continue climate adaptation and
            conservation efforts.
            """

        doc.add_heading(
            "Recommendations",
            level=2
        )

        doc.add_paragraph(
            recommendation
        )
        # ==============================
        # CONCLUSION
        # ==============================

        if mean_lst > 40 and mean_ndvi < 0.2:

            conclusion = f"""
            The assessment revealed severe thermal
            conditions with temperatures ranging from
            {lst_min:.2f} °C to {lst_max:.2f} °C and
            an average temperature of {mean_lst:.2f} °C.

            The low NDVI value of {mean_ndvi:.2f}
            indicates limited vegetation cover, which
            is likely contributing to elevated urban
            heat intensity.

            Immediate climate adaptation and urban
            greening interventions are recommended to
            improve environmental resilience.
            """

        elif mean_lst > 35 and mean_ndvi < 0.5:

            conclusion = f"""
            The assessment recorded temperatures ranging
            from {lst_min:.2f} °C to {lst_max:.2f} °C,
            with an average temperature of {mean_lst:.2f} °C.

            The mean NDVI value of {mean_ndvi:.2f}
            indicates moderate vegetation conditions
            across the study area.

            Strategic environmental management and
            targeted greening initiatives could help
            reduce future urban heat risks.
            """

        else:

            conclusion = f"""
            The assessment recorded temperatures ranging
            from {lst_min:.2f} °C to {lst_max:.2f} °C,
            with an average temperature of {mean_lst:.2f} °C.

            The vegetation condition represented by
            a mean NDVI value of {mean_ndvi:.2f}
            contributes positively to environmental
            quality and thermal regulation.

            Continued sustainable land management is
            recommended to maintain climate resilience
            and environmental sustainability.
            """

        doc.add_heading(
            "Conclusion",
            level=2
        )

        doc.add_paragraph(
            conclusion
        )
        # ==============================
        # SAVE DOCUMENT
        # ==============================

        doc.save(
            report_path
        )
        with open(
            report_path,
            "rb"
        ) as doc_file:

            st.download_button(
                label="Download Word Report",
                data=doc_file,
                file_name="Urban_Climate_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        st.success(
            "Word report generated successfully!"
        )